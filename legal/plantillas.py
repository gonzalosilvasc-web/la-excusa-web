"""Módulo 2: biblioteca de modelos Word y detección de placeholders."""

from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

import pandas as pd
from docx import Document

from .config import (
    BASE_DIR,
    PLACEHOLDERS_ESTANDAR,
    TEMPLATES_DIR,
    TEMPLATES_XLSX,
)
from .storage import append_row, read_excel_or_empty
from .textutil import sanitizar_nombre_archivo, uid

_RE_PLACEHOLDER = re.compile(r"\{\{\s*([\w\.]+)\s*\}\}")


def texto_completo_docx(ruta: Path) -> str:
    """Extrae todo el texto de un .docx (párrafos, tablas, headers y footers)."""
    doc = Document(str(ruta))
    partes: list[str] = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                partes.append(celda.text)
    for seccion in doc.sections:
        for cont in (seccion.header, seccion.footer):
            partes.extend(p.text for p in cont.paragraphs)
    return "\n".join(partes)


def detectar_placeholders(ruta: Path) -> list[str]:
    """Detecta placeholders {{VARIABLE}} en un .docx.

    Combina docxtpl (que ve el XML crudo, robusto ante 'runs' partidos) con un
    regex sobre el texto plano como respaldo.
    """
    encontrados: set[str] = set()
    try:
        from docxtpl import DocxTemplate

        tpl = DocxTemplate(str(ruta))
        encontrados |= set(tpl.get_undeclared_template_variables())
    except Exception:
        pass
    try:
        encontrados |= set(_RE_PLACEHOLDER.findall(texto_completo_docx(ruta)))
    except Exception:
        pass
    return sorted(encontrados)


def cargar_templates(solo_activos: bool = False) -> pd.DataFrame:
    df = read_excel_or_empty(TEMPLATES_XLSX)
    if solo_activos and not df.empty:
        df = df[df["Activo"].astype(str).str.lower().isin(("true", "1", "si", "sí"))]
    return df


def guardar_template(
    archivo_origen: Path,
    metadatos: dict[str, Any],
) -> dict[str, Any]:
    """Guarda un modelo Word en database/templates/ y lo indexa."""
    id_template = uid("TPL")
    placeholders = detectar_placeholders(archivo_origen)
    nombre = sanitizar_nombre_archivo(
        f"{metadatos.get('Nombre_Template', 'modelo')}_{id_template.split('-')[-1]}"
    ) + ".docx"
    destino = TEMPLATES_DIR / nombre
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(archivo_origen, destino)

    registro = {
        "ID_Template": id_template,
        "Nombre_Template": str(metadatos.get("Nombre_Template", "")).strip(),
        "Tipo_Demanda": str(metadatos.get("Tipo_Demanda", "")).strip(),
        "Banco_Entidad": str(metadatos.get("Banco_Entidad", "")).strip(),
        "Moneda": str(metadatos.get("Moneda", "")).strip(),
        "Numero_Titulos": str(metadatos.get("Numero_Titulos", "1")),
        "Tiene_Hipoteca": bool(metadatos.get("Tiene_Hipoteca", False)),
        "Tiene_Prenda": bool(metadatos.get("Tiene_Prenda", False)),
        "Tiene_Avalistas": bool(metadatos.get("Tiene_Avalistas", False)),
        "Tiene_Codeudores": bool(metadatos.get("Tiene_Codeudores", False)),
        "Tipo_Instrumento": str(metadatos.get("Tipo_Instrumento", "")).strip(),
        "Ruta_Template": str(destino.relative_to(BASE_DIR)),
        "Placeholders_Detectados": ", ".join(placeholders),
        "Fecha_Carga": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "Usuario_Carga": str(metadatos.get("Usuario_Carga", "")).strip() or "anonimo",
        "Observaciones": str(metadatos.get("Observaciones", "")).strip(),
        # Sin placeholders no se permite generación automática directa
        "Activo": bool(metadatos.get("Activo", True)) and bool(placeholders),
    }
    append_row(TEMPLATES_XLSX, registro)
    return registro


def crear_plantilla_base(nombre_archivo: str = "modelo_base_demanda_ejecutiva.docx") -> Path:
    """Crea un modelo Word base con todos los placeholders estándar.

    Sirve como plantilla funcional de fábrica y como ejemplo de estructura
    para los modelos propios del estudio.
    """
    destino = TEMPLATES_DIR / nombre_archivo
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    p = doc.add_paragraph()
    p.add_run(
        "EN LO PRINCIPAL: DEMANDA EJECUTIVA; PRIMER OTROSÍ: ACOMPAÑA DOCUMENTOS "
        "CON CITACIÓN; SEGUNDO OTROSÍ: SEÑALA BIENES PARA LA TRABA DE EMBARGO; "
        "TERCER OTROSÍ: PERSONERÍA; CUARTO OTROSÍ: PATROCINIO Y PODER; "
        "QUINTO OTROSÍ: FORMA DE NOTIFICACIÓN."
    ).bold = True
    doc.add_paragraph("")
    doc.add_paragraph("{{TRIBUNAL}}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "{{REPRESENTANTE_DEMANDANTE}}, abogado, en representación convencional de "
        "{{DEMANDANTE}}, RUT {{RUT_DEMANDANTE}}, según se acreditará, a US. "
        "respetuosamente digo:"
    )
    doc.add_paragraph(
        "Que, en la representación que invisto, vengo en deducir demanda "
        "ejecutiva de {{MATERIA}}, en procedimiento {{PROCEDIMIENTO}}, en "
        "contra de las personas que a continuación se individualizan:"
    )
    doc.add_paragraph("{{DEMANDADOS}}")
    doc.add_paragraph("")
    doc.add_paragraph("I. LOS TÍTULOS EJECUTIVOS").runs[0].bold = True
    doc.add_paragraph("{{TITULOS_EJECUTIVOS}}")
    doc.add_paragraph("")
    doc.add_paragraph("II. RELACIÓN DE LOS HECHOS Y EXIGIBILIDAD").runs[0].bold = True
    doc.add_paragraph("{{RELATO_DE_LOS_HECHOS}}")
    doc.add_paragraph(
        "El capital total adeudado asciende a {{MONTO_CAPITAL}} ({{MONEDA}})."
        " {{EQUIVALENCIA_PESOS}}"
    )
    doc.add_paragraph("")
    doc.add_paragraph("POR TANTO,").runs[0].bold = True
    doc.add_paragraph("{{PETITORIO}}")
    doc.add_paragraph("")
    doc.add_paragraph("PRIMER OTROSÍ: ").runs[0].bold = True
    doc.add_paragraph(
        "Ruego a US. tener por acompañados, con citación, los siguientes documentos:"
    )
    doc.add_paragraph("{{DOCUMENTOS_PRIMER_OTROSI}}")
    doc.add_paragraph("")
    doc.add_paragraph("SEGUNDO OTROSÍ: ").runs[0].bold = True
    doc.add_paragraph(
        "Ruego a US. tener presente que, para la traba del embargo, señalo los "
        "siguientes bienes de propiedad de la parte demandada: {{BIENES_EMBARGO}}"
    )
    doc.add_paragraph("")
    doc.add_paragraph("TERCER OTROSÍ: ").runs[0].bold = True
    doc.add_paragraph("Ruego a US. tener presente lo siguiente: {{TEXTO_PERSONERIA}}")
    doc.add_paragraph("")
    doc.add_paragraph("CUARTO OTROSÍ: ").runs[0].bold = True
    doc.add_paragraph(
        "Ruego a US. tener presente que, en mi calidad de abogado habilitado "
        "para el ejercicio de la profesión, asumo personalmente el patrocinio y "
        "poder de la presente causa."
    )
    doc.add_paragraph("")
    doc.add_paragraph("QUINTO OTROSÍ: ").runs[0].bold = True
    doc.add_paragraph(
        "Ruego a US. tener presente que solicito que las resoluciones que se "
        "dicten en autos me sean notificadas electrónicamente."
    )
    doc.add_paragraph("")
    doc.add_paragraph("{{GARANTIAS}}")
    doc.add_paragraph("")
    doc.add_paragraph("Fecha de la presentación: {{FECHA_DEMANDA}}")

    doc.save(str(destino))
    return destino


def registrar_plantilla_base_si_falta() -> Optional[dict[str, Any]]:
    """Registra la plantilla base de fábrica si la biblioteca está vacía."""
    df = cargar_templates()
    if not df.empty:
        return None
    ruta = crear_plantilla_base()
    return guardar_template(
        ruta,
        {
            "Nombre_Template": "Modelo base demanda ejecutiva (fábrica)",
            "Tipo_Demanda": "Cobro de pagaré en pesos",
            "Banco_Entidad": "Genérico",
            "Moneda": "CLP",
            "Numero_Titulos": "1+",
            "Tipo_Instrumento": "Pagare",
            "Usuario_Carga": "sistema",
            "Observaciones": "Plantilla base generada automáticamente. "
            "Reemplácela por los modelos propios del estudio.",
            "Activo": True,
        },
    )


def opciones_sin_placeholders() -> list[str]:
    """Opciones ofrecidas cuando un modelo no tiene placeholders."""
    return [
        "Usarlo solo como referencia de estilo (no generación automática)",
        "Crear una copia con placeholders sugeridos",
        "Mapear variables manualmente más adelante",
    ]


def crear_copia_con_placeholders(ruta_original: Path) -> Path:
    """Crea una copia del modelo agregándole la estructura estándar de placeholders."""
    doc = Document(str(ruta_original))
    doc.add_page_break()
    doc.add_paragraph(
        "=== SECCIÓN AGREGADA AUTOMÁTICAMENTE: PLACEHOLDERS SUGERIDOS ==="
    ).runs[0].bold = True
    doc.add_paragraph(
        "Copie estos placeholders dentro del cuerpo del modelo donde corresponda "
        "y elimine esta sección:"
    )
    for ph in PLACEHOLDERS_ESTANDAR:
        doc.add_paragraph("{{" + ph + "}}")
    destino = ruta_original.with_name(ruta_original.stem + "_con_placeholders.docx")
    doc.save(str(destino))
    return destino
